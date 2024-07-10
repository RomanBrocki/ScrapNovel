from pattern import pattern
from pattern import replacement
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
import os
from shutil import copyfile
from docx import Document


class ClasseEbook:
    """
    Bibliotecas necessárias: selenium, selenium.webdriver.common.by ; selenium.webdriver.chrome.service;
    webdriver-manager;shutil; python-docx.
    Utiliza 1 arquivo assistente pattern, com lista de possiveis ocorrências de marcas dágua anti-srapping
    Classe que recebe dados para geração do scrap, e possui funções para correr o scrap e gerar ebook.

    Atributos:
    nome = Nome do ebook recebido
    site_inicio = site de inicio do scrap
    site_final = site final do scrap
    lista_ocorrencias = lista construida durante o scrap com ocorrências identificadas do pattern
    ocorrencias = valor inteiro começando em 0, agregando instancia encontrada de ocorrência
    caminho = Local de destino do ebook. Começa vazio para receber posteriormente o arquivo
    profile = Caminho do profile do chrome a ser usado pelo navegador. De preferência com adblocker instalado.

    def copiarcapitulos(self, site_inicio, site_final):
        recebe info do primeiro e último capítulos,
        usa o chrome para fazer o scrap
        processa o scrap gerando arquivo Novel.docx (ebook) na pasta raiz e log.docx (log ocorrências) na pasta raiz.
        ambos docx serão apagados após acionamento das funções passarArquivo e pegaLog

    def pegaLog(self):
        processa o arquivo de log criando e retornando variável com seu conteúdo. Apaga o arquivo log.docx

    def passarArquivo(self, caminho, nome_ebook)
        processa o arquivo gerado pelo copiarcapitulos
        Gera arquivo ebook com nome definido na pasta de caminho definido
        apaga arquivo Novel.docx

    def scrapNovel(self, nome_ebook, caminho, site_inicio, site_final):
        Função base da Classe para uso dos seus recursos por meio de código direto.
    """

    def __init__(self):
        self.nome_ebook = ""
        self.site_inicio = ""
        self.site_final = ""
        self.lista_ocorrencias = []
        self.ocorrencias = 0
        self.caminho = ""
        self.profile = r"c:\Users\Roman\AppData\Local\Google\Chrome\User Data\Profile Selenium"

    def copiarcapitulos(self, site_inicio, site_final):
        self.site_inicio = site_inicio
        self.site_final = site_final
        options = webdriver.ChromeOptions()
        options.add_argument("user-data-dir=" + self.profile)
        #servico = Service(r'F:\OneDrive\Documentos\Python\chromedriver.exe')
        servico = Service(ChromeDriverManager().install())
        navegador = webdriver.Chrome(service=servico, options=options)
        documento = Document()
        documento.save('Novel.docx')
        while True:
            navegador.get(self.site_inicio)
            titulo = navegador.find_element(By.CLASS_NAME, 'chr-text')
            conteudo = navegador.find_element(By.ID, 'chr-content')
            documento = Document("Novel.docx")
            capitulo = conteudo.text
            for elemento in pattern:
                if elemento in capitulo:
                    self.lista_ocorrencias.append(f'Encontrado:{elemento}')
                    self.ocorrencias += 1
                conteudo_tratado = capitulo.replace(elemento, replacement)
                capitulo = conteudo_tratado
            paragrafo1 = documento.add_paragraph(titulo.text)
            paragrafo2 = documento.add_paragraph(capitulo)
            documento.add_page_break()
            documento.save("Novel.docx")
            if self.site_inicio == self.site_final:
                break
            links = navegador.find_elements(By.ID, 'next_chap')
            listalinks = []
            for lnk in links:
                listalinks.append(lnk.get_attribute('href'))
            self.site_inicio = listalinks[0]
        log = Document()
        for ocorrencia in self.lista_ocorrencias:
            paragrafo = log.add_paragraph(ocorrencia)
        total = log.add_paragraph(f'A quantidade total de ocorrências foi {self.ocorrencias}')
        log.save('log.docx')

        print(self.ocorrencias)
        navegador.quit()

    def pegaLog(self):
        list_log = ['Segue Lista de ocorrências encontradas:']
        document = Document('log.docx')
        for p in document.paragraphs:
            list_log.append(p.text)
        os.remove('log.docx')
        return list_log

    def passarArquivo(self, caminho, nome_ebook):
        self.nome_ebook = nome_ebook
        self.caminho = caminho
        documento = Document()
        copyfile("Novel.docx", r"{}/{}.docx".format(self.caminho, self.nome_ebook))
        os.remove('Novel.docx')

    def scrapNovel(self, nome_ebook, caminho, site_inicio, site_final):
        self.copiarcapitulos(site_inicio, site_final)
        self.passarArquivo(caminho, nome_ebook)
